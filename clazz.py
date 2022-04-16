from pathlib import Path
from regex import *
from extractor import *
from setup import *
from docx.table import Table
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from table_worker import create_table, fill_table, create_table_for_class_members


class Class:

    def __init__(self, class_code: string):
        self.comment = ""
        self.name = ""
        self.methods = []
        self.params = {}
        self.properties = []
        self.__extract_class_info(class_code)
        self.methods.sort(key=lambda meth: meth.name)
        self.properties.sort(key=lambda pr: pr.name)

    def __extract_class_info(self, code: string):
        self.__extract_class_name(code)
        self.__extract_comment(code)

    def __extract_class_name(self, info: string):
        regex = rf'{class_word} [\w\d]*(<([\w\d ]*,?)+>)?'
        matches = re.search(regex, info)
        if matches:
            self.name = matches.group().replace(class_word, "").strip()
            print(f'{info} -> {self.name}')
        else:
            raise Exception("Incorrect file. Cannot extract class name")

    def __extract_comment(self, info: string):
        matches = re.search(comment_regex, info, re.MULTILINE)
        if matches:
            self.comment = extract_comment(matches.group())
            self.params = extract_params(matches.group())

    def to_row(self, table: Table):
        cells = table.add_row().cells
        cells[0].text = self.name
        cells[1].text = self.comment
        cells[0].width = Pt(100)
        cells[1].width = Pt(100)
        for i in range(2):
            cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            cells[i].paragraphs[0].runs[0].font.size = Pt(12)

    def class_to_table(self, doc: Document(), number: int):
        if len(self.methods) > 0 or len(self.properties) > 0:
            if len(self.methods) > 0:
                headers = ['Имя', 'Модификатор', "Тип", "Параметры", "Назначение"]
                table_name = f'Таблица 2.{number}. Описание методов класса {self.name}'
                create_table_for_class_members(doc, table_name, headers, self.methods)

            num = number
            if len(self.properties) > 0:
                if len(self.methods) > 0:
                    doc.add_paragraph()
                    num += 1
                headers = ['Имя', 'Модификатор', "Тип", "Назначение"]
                table_name = f'Таблица 2.{num}. Описание свойств класса {self.name}'
                create_table_for_class_members(doc, table_name, headers, self.properties)

            return num
        else:
            return number