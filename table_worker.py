import string
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.table import Table


def create_table_for_class_members(doc: Document, table_name: string, headers: [], entities: []):
    create_table_header(doc, table_name)
    table = create_table(doc, headers)
    fill_table(table, entities)


def create_table_header(doc: Document, header: string):
    p = doc.add_paragraph(header)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(12)


def create_table(doc: Document, headers: []):
    table = doc.add_table(1, len(headers))
    table.style = 'Table Grid'
    table.autofit = True
    head_cells = table.rows[0].cells
    for i, item in enumerate(headers):
        p = head_cells[i].paragraphs[0]
        p.add_run(item).bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.runs[0].font.name = 'Times New Roman'
    return table


def fill_table(table: Table, entities: list):
    for entity in entities:
        entity.to_row(table)
