import re
from pathlib import Path
import string

import setup
from regex import *
from clazz import Class
from interface import Interface
from method import Method
from property import Property
from docx import Document
from table_worker import create_table_for_class_members


class Converter:

    def __init__(self, path: string):
        self.path = Path(path)
        self.classes = []
        self.interfaces = []
        self.__read_files(self.path)
        self.doc = Document()
        self.interfaces.sort(key=lambda elm: elm.name)
        self.classes.sort(key=lambda elm: elm.name)


    def __read_files(self, path: Path):
        if path.is_dir():
            for subdir in path.iterdir():
                self.__read_files(subdir)
        else:
            if path.name.endswith(setup.file_extension):
                self.__read_file(path)

    def __read_file(self, path: Path):
        if not path.is_file():
            raise Exception("Path have to be a file")
        code = path.read_text()
        regex_for_class = get_regex_for(class_regex)
        class_matches = re.findall(regex_for_class, code, re.MULTILINE)
        for class_code in class_matches:
            c = Class(class_code[0])
            regex_for_method = get_regex_for(method_regex)
            methods = re.findall(regex_for_method, code, re.MULTILINE)
            for meth in methods:
                m = Method(meth[0])
                c.methods.append(m)

            regex_for_property = get_regex_for(property_regex)
            properties = re.findall(regex_for_property, code, re.MULTILINE)
            for pr in properties:
                p = Property(pr[0])
                c.properties.append(p)

            self.classes.append(c)

        regex_for_interface = get_regex_for(interface_regex)
        interface_matches = re.findall(regex_for_interface, code, re.MULTILINE)
        for interface_code in interface_matches:
            inter = Interface(interface_code[0])

            regex_for_method = get_regex_for(method_regex)
            methods = re.findall(regex_for_method, code, re.MULTILINE)
            for meth in methods:
                m = Method(meth[0])
                inter.methods.append(m)

            regex_for_property = get_regex_for(property_regex)
            properties = re.findall(regex_for_property, code, re.MULTILINE)
            for pr in properties:
                p = Property(pr[0])
                inter.properties.append(p)

            self.interfaces.append(inter)

    def crete_class_table(self):
        headers = ['Интерфейс', 'Назначение']
        table_name = 'Таблица 1.1. Классы'
        create_table_for_class_members(self.doc, table_name, headers, self.classes)

    def crete_class_members_tables(self):
        self.classes.sort(key=lambda elm: elm.name)
        number = 1
        for c in self.classes:
            self.doc.add_paragraph()
            number = c.class_to_table(self.doc, number) + 1

    def crete_interface_table(self):
        p = self.doc.add_paragraph()
        headers = ['Интерфейс', 'Назначение']
        table_name = 'Таблица 1.2. Интерфейсы'
        create_table_for_class_members(self.doc, table_name, headers, self.interfaces)

    def crete_interface_members_tables(self):
        self.interfaces.sort(key=lambda elm: elm.name)
        number = 1
        for c in self.interfaces:
            self.doc.add_paragraph()
            number = c.interface_to_table(self.doc, number) + 1

    def save_table(self, path: string):
        self.doc.save(path)
